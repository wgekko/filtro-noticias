from datetime import datetime
import streamlit as st
import feedparser
import pandas as pd
import requests
import time
import json
import io
from docx import Document
from textblob import TextBlob  # pip install textblob
import smtplib
import ssl
import altair as alt
import pytz

# ----- CONFIG -----
OLLAMA = "http://localhost:11434"
MODEL = "deepseek-r1:14b"

RSS_MAP = {
    "Infobae": "https://www.infobae.com/feeds/rss/economia.xml",
    "AmbitoFin": "https://www.ambito.com/rss/finanzas.xml",
    "AmbitoEcon": "https://www.ambito.com/rss/economia.xml",
    "CronistaFin": "https://www.cronista.com/files/rss/finanzas-mercados.xml",
    "CronistaEcon": "https://www.cronista.com/files/rss/economia-politica.xml",
    "LaNacion": "https://www.lanacion.com.ar/rss/economia.xml",
    "Clarin":  "https://www.clarin.com/rss/economia",        
}

# ----- FUNCIONES -----

def ollama_tldr(text: str) -> str:
    try:
        r = requests.post(f"{OLLAMA}/api/generate",
            json={"model": MODEL,
                  "prompt": f"Genera un resumen de una oración del siguiente contenido de noticias e imprima solo esa oración.:\n{text}",
                  "temperature": 0.3,
                  "stream": False},
            timeout=60)
        r.raise_for_status()
        return r.json()["response"].strip()
    except Exception as e:
        return f"(Resumen no disponible: {e})"

def analyze_sentiment(text: str) -> str:
    analysis = TextBlob(text)
    if analysis.sentiment.polarity > 0.1:
        return "Positiva"
    elif analysis.sentiment.polarity < -0.1:
        return "Negativa"
    else:
        return "Neutra"

@st.cache_data(ttl=1800)
def fetch_papers(feeds, total_max_n):
    rows = []
    per_source = max(5, total_max_n // len(feeds))
    argentina_tz = pytz.timezone('America/Argentina/Buenos_Aires') # Definir aquí o pasar como argumento

    for tag in feeds:
        parsed = feedparser.parse(RSS_MAP[tag])
        if not parsed.entries:
            # Opcional: informar al usuario que un feed está vacío o no accesible
            # st.info(f"No se encontraron entradas recientes para {tag} o el feed no está accesible.")
            pass # Puedes decidir si quieres mostrar un mensaje

        count = 0
        for e in parsed.entries:
            if count >= per_source:
                break

            summary = e.summary.replace("\n", " ") if 'summary' in e else ""
            tldr = ollama_tldr(summary) # Asumo que quieres el resumen del 'summary' original
            sentiment = analyze_sentiment(tldr) # Analizar sentimiento del resumen corto

            pub_date_final_str = "" # Valor por defecto
            if 'published_parsed' in e and e.published_parsed:
                try:
                    # Crear un datetime naive a partir de la tupla de feedparser
                    dt_naive = datetime.fromtimestamp(time.mktime(e.published_parsed))
                    # Localizar a UTC (asumiendo que las fechas de feedparser pueden ser UTC o sin tz)
                    # y luego convertir a Argentina
                    dt_argentina = pytz.utc.localize(dt_naive).astimezone(argentina_tz)
                    #pub_date_final_str = dt_argentina.strftime("%Y-%m-%d %H:%M:%S")                    
                    pub_date_final_str = dt_argentina.strftime("%d-%m-%Y %H:%M")
                except Exception as ex:
                    # Si falla la conversión detallada, usar solo la fecha como antes
                    pub_date_final_str = e.published[:10] if 'published' in e else "Fecha no disponible"
                    # st.warning(f"Error convirtiendo fecha para {e.title}: {ex}")
            elif 'published' in e:
                 pub_date_final_str = e.published[:10] # Fallback a solo fecha si no hay 'published_parsed'
            else:
                pub_date_final_str = "Fecha no disponible"
            rows.append({
                "source": tag,
                "date": pub_date_final_str,
                "title": e.title,
                "link": e.link,
                "tldr": tldr,
                "sentiment": sentiment
            })
            count += 1
            time.sleep(0.1) # Mantener para no sobrecargar los servidores RSS
    return pd.DataFrame(rows)

def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Noticias')
    return output.getvalue()

def to_word(df):
    doc = Document()
    doc.add_heading('Noticias de Economía', 0)
    for _, row in df.iterrows():
        doc.add_heading(row["title"], level=1)
        doc.add_paragraph(f"Fuente: {row['source']}")
        doc.add_paragraph(f"Fecha: {row['date']}")
        doc.add_paragraph(f"Resumen: {row['tldr']}")
        doc.add_paragraph(f"Sentimiento: {row['sentiment']}")
        doc.add_paragraph(f"Enlace: {row['link']}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def send_email(subject, body):
    try:
        sender = "tucorreo@gmail.com"
        receiver = "destinatario@gmail.com"
        password = "tu_clave_de_app"
        message = f"Subject: {subject}\n\n{body}"

        context = ssl.create_default_context()
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
            server.login(sender, password)
            server.sendmail(sender, receiver, message)
        st.success("📧 Notificación enviada por correo.")
    except Exception as e:
        st.error(f"Error enviando email: {e}")

def send_telegram(message):
    try:
        bot_token = "TU_BOT_TOKEN"
        chat_id = "TU_CHAT_ID"
        url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
        params = {"chat_id": chat_id, "text": message}
        response = requests.get(url, params=params)
        if response.ok:
            st.success("Mensaje enviado por Telegram.", icon=":material/done_all:")
        else:
            st.error("Falló el envío por Telegram.", icon=":material/error:")
    except Exception as e:
        st.error(f"Error en Telegram: {e}")

# ----- UI -----
st.set_page_config(page_title="Noticias Económicas", layout="wide", page_icon=":material/newspaper:")
with open("asset/styles.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

st.markdown("<h1>📈 Resumen de Información Económica Inteligente</h1>" ,unsafe_allow_html=True)

st.sidebar.title("Opciones")
sel = st.sidebar.multiselect("📰 Fuentes", list(RSS_MAP), ["Infobae", "AmbitoFin", "AmbitoEcon", "CronistaFin", "CronistaEcon", "LaNacion", "Clarin"])
max_n = st.sidebar.slider("📑 Noticias por fuente", 5, 30, 10)

if st.sidebar.button("Actualizar", icon=":material/autorenew:"):
    st.cache_data.clear()

df = fetch_papers(tuple(sel), max_n)

#query = st.text_input("🔍 Buscar por palabra clave")
#if query:
#    df = df[df["tldr"].str.contains(query, case=False) | df["title"].str.contains(query, case=False)]
query = st.text_input("🔍 Buscar por palabra clave")
if query:
    df = df[df["tldr"].fillna('').str.contains(query, case=False) | df["title"].fillna('').str.contains(query, case=False)]

if not df.empty:
    #st.write("###")
    st.write("---")
    with st.container(border=True):
        # Reordenamos las columnas
        ordered_cols = ["source", "date", "title", "sentiment", "tldr", "link"]
        df = df[[col for col in ordered_cols if col in df.columns]]        
        #st.subheader("📋 Tabla de Noticias")
        #st.dataframe(df, use_container_width=True)    
    with st.container(border=True):        
        st.dataframe(df, use_container_width=True)
        st.subheader("Descargar la información en distintos formatos")
        col1, col2, col3 = st.columns(3, vertical_alignment="center")
        with col1:
            csv = df.to_csv(index=False, encoding="utf-8-sig")
            st.download_button("Descargar CSV", csv, "noticias.csv", mime="text/csv", icon=":material/download:")
        with col2:
            st.download_button("Descargar Excel", to_excel(df), "noticias.xlsx", icon=":material/download:" )        
        with col3:        
            st.download_button("Descargar Word", to_word(df), "noticias.docx", icon=":material/download:" )     
    
    
    with st.container(border=True):   
        st.subheader("Enviar la información por mail o telegram")
        col3, col4 = st.columns(2, vertical_alignment="center")
      
        with col3:
            if st.button("Enviar notificación por correo", icon=":material/local_post_office:"):
                send_email("Noticias Económicas", df.to_string())    
        with col4:             
            if st.button("Enviar notificación por Telegram",  icon=":material/send:"):
                send_telegram("Nuevas noticias disponibles. Revisá Radar Económico.")
                
    with st.container(border=True):                
        st.subheader("Sentimiento por Fuente")
        if not df.empty:
            import altair as alt

            # Agrupamos los datos
            chart_data = df.groupby(["source", "sentiment"]).size().reset_index(name="count")

            # Creamos un gráfico de barras apiladas
            chart = alt.Chart(chart_data).mark_bar().encode(
                x=alt.X('source:N', title='Fuente', axis=alt.Axis(labelAngle=0)),
                y=alt.Y('count:Q', title='Cantidad'),
                color=alt.Color('sentiment:N', title="Sentimiento", scale=alt.Scale(
                    domain=["Positiva", "Neutra", "Negativa"],
                    range=["#2ecc71", "#f1c40f", "#e74c3c"]
                )),
                tooltip=["source", "sentiment", "count"]
            ).properties(
                height=400
            )

            # Mostrar el gráfico ajustado al contenedor
            st.altair_chart(chart, use_container_width=True)

    st.caption("Resúmenes con IA local (Gemma 3:4B en Ollama)")
else:
    st.warning("No se encontraron noticias para mostrar.", icon=":material/warning:")


# --------------- footer -----------------------------
st.write("---")
with st.container():
  #st.write("---")
  st.write("&copy; - derechos reservados -  2025 -  Walter Gómez - FullStack Developer - Data Science - Business Intelligence")
  #st.write("##")
  left, right = st.columns(2, gap='medium', vertical_alignment="bottom")
  with left:
    #st.write('##')
    st.link_button("Mi LinkedIn", "https://www.linkedin.com/in/walter-gomez-fullstack-developer-datascience-businessintelligence-finanzas-python/",use_container_width=True)
  with right: 
     #st.write('##') 
    st.link_button("Mi Porfolio", "https://walter-portfolio-animado.netlify.app/", use_container_width=True)
      