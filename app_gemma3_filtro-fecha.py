from datetime import datetime, timedelta
import streamlit as st
import feedparser
import pandas as pd
import requests
import time
import json
import io
from docx import Document
from textblob import TextBlob  # pip install textblob
import smtplib
import ssl
import altair as alt
import pytz

from app_modelo_gemma3 import send_email, send_telegram

# ----- CONFIG ----- 
OLLAMA = "http://localhost:11434"
MODEL = "gemma3:latest"

RSS_MAP = {
    "Infobae": "https://www.infobae.com/feeds/rss/economia.xml",
    "AmbitoFin": "https://www.ambito.com/rss/finanzas.xml",
    "AmbitoEcon": "https://www.ambito.com/rss/economia.xml",
    "CronistaFin": "https://www.cronista.com/files/rss/finanzas-mercados.xml",
    "CronistaEcon": "https://www.cronista.com/files/rss/economia-politica.xml",
    "LaNacion": "https://www.lanacion.com.ar/rss/economia.xml",
    "Clarin":  "https://www.clarin.com/rss/economia",        
}

argentina_tz = pytz.timezone('America/Argentina/Buenos_Aires')

# ----- FUNCIONES ----- 
def ollama_tldr(text: str) -> str:
    try:
        r = requests.post(f"{OLLAMA}/api/generate",
            json={"model": MODEL,
                  "prompt": f"Genera un resumen de una oración del siguiente contenido de noticias e imprima solo esa oración.:\n{text}",
                  "temperature": 0.3,
                  "stream": False},
            timeout=60)
        r.raise_for_status()
        return r.json()["response"].strip()
    except Exception as e:
        return f"(Resumen no disponible: {e})"

def analyze_sentiment(text: str) -> str:
    analysis = TextBlob(text)
    if analysis.sentiment.polarity > 0.1:
        return "Positiva"
    elif analysis.sentiment.polarity < -0.1:
        return "Negativa"
    else:
        return "Neutra"

@st.cache_data(ttl=3600)
def fetch_papers(feeds, total_max_n):
    rows = []
    per_source = max(5, total_max_n // len(feeds))
    
    for tag in feeds:
        parsed = feedparser.parse(RSS_MAP[tag])
        count = 0
        for e in parsed.entries:
            if count >= per_source:
                break
            summary = e.summary.replace("\n", " ") if 'summary' in e else ""
            tldr = ollama_tldr(summary)
            sentiment = analyze_sentiment(tldr)

            pub_date = None
            if 'published_parsed' in e and e.published_parsed:
                pub_date = datetime(*e.published_parsed[:6])
                pub_date = pub_date.astimezone(argentina_tz)

            rows.append({
                "source": tag,
                "date": pub_date,
                "title": e.title,
                "link": e.link,
                "tldr": tldr,
                "sentiment": sentiment
            })
            count += 1
            time.sleep(0.1)
    return pd.DataFrame(rows)

# Cambiar la columna "date" a formato "%d-%m-%Y" y eliminar la zona horaria
def format_date(df):
    if "date" in df.columns:
        # Convertir a datetime si no lo está
        df["date"] = pd.to_datetime(df["date"], errors='coerce')  # Convertir a datetime, ignorando errores
        
        # Eliminar la zona horaria si está presente
        df["date"] = df["date"].dt.tz_localize(None) if df["date"].dt.tz is not None else df["date"]

        # Formatear la columna "date" al formato "dd-mm-yyyy"
        df["date"] = df["date"].dt.strftime('%d-%m-%Y')
    return df

# Función para convertir el DataFrame a un archivo Excel
def to_excel(df):
    # Asegurarse de que "date" esté correctamente convertido a datetime
    df = format_date(df)  # Aseguramos la conversión a datetime y el formato adecuado

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Noticias')
    return output.getvalue()

# Función para convertir el DataFrame a un archivo Word
def to_word(df):
    doc = Document()
    doc.add_heading('Noticias de Economía', 0)
    for _, row in df.iterrows():
        doc.add_heading(row["title"], level=1)
        doc.add_paragraph(f"Fuente: {row['source']}")
        doc.add_paragraph(f"Fecha: {row['date']}")
        doc.add_paragraph(f"Resumen: {row['tldr']}")
        doc.add_paragraph(f"Sentimiento: {row['sentiment']}")
        doc.add_paragraph(f"Enlace: {row['link']}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ----- UI ----- 
st.set_page_config(page_title="Noticias Económicas", layout="wide", page_icon=":material/newspaper:")

with open("asset/styles.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

st.markdown("<h1>📈 Resumen de Información Económica Inteligente</h1>", unsafe_allow_html=True)

st.sidebar.title("Opciones")
sel = st.sidebar.multiselect("📰 Fuentes", list(RSS_MAP), list(RSS_MAP.keys()))
max_n = st.sidebar.slider("📑 Noticias por fuente", 5, 30, 10)

# Filtro por rango de días
st.sidebar.markdown("### 📅 Rango de fechas")
rango_dias = st.sidebar.selectbox("Seleccionar rango", options=[1, 3, 5], index=1, format_func=lambda x: f"Últimos {x} días")
hoy = datetime.now().date()
fecha_inicio = hoy - timedelta(days=rango_dias)
fecha_fin = hoy

if st.sidebar.button("Actualizar", icon=":material/autorenew:"):
    st.cache_data.clear()

df = fetch_papers(tuple(sel), max_n)

# Filtrar por fechas
df = df[df["date"].notnull()]
df["date"] = pd.to_datetime(df["date"], format='%d-%m-%Y', errors='coerce')  # Convertir de nuevo a datetime

# Asegúrate de que `fecha_inicio` y `fecha_fin` sean objetos `datetime.date`
df = df[df["date"].dt.date.between(fecha_inicio, fecha_fin)]

# Aplicar formato a la columna "date"
df = format_date(df)

query = st.text_input("🔍 Buscar por palabra clave")
if query:
    df = df[df["tldr"].fillna('').str.contains(query, case=False) | df["title"].fillna('').str.contains(query, case=False)]

if not df.empty:
    st.write("---")
    with st.container(border=True):
        ordered_cols = ["source", "date", "title", "sentiment", "tldr", "link"]
        df = df[[col for col in ordered_cols if col in df.columns]]

    with st.container(border=True):        
        st.dataframe(df, use_container_width=True)
        st.subheader("Descargar la información en distintos formatos")
        col1, col2, col3 = st.columns(3)
        with col1:
            csv = df.to_csv(index=False, encoding="utf-8-sig")
            st.download_button("CSV", csv, "noticias.csv", mime="text/csv", icon=":material/download:")
        with col2:
            st.download_button("Excel", to_excel(df), "noticias.xlsx", icon=":material/download:")        
        with col3:        
            st.download_button("Word", to_word(df), "noticias.docx", icon=":material/download:")     

    with st.container(border=True):   
        st.subheader("Enviar la información por mail o telegram")
        col3, col4 = st.columns(2)
        with col3:
            if st.button("Enviar notificación por correo", icon=":material/local_post_office:"):
                send_email("Noticias Económicas", df.to_string())    
        with col4:             
            if st.button("Enviar notificación por Telegram", icon=":material/send:"):
                send_telegram("Nuevas noticias disponibles. Revisá Radar Económico.")
                
    with st.container(border=True):
        st.subheader("Sentimiento por Fuente")
        if not df.empty:
            chart_data = df.groupby(["source", "sentiment"]).size().reset_index(name="count")
            chart = alt.Chart(chart_data).mark_bar().encode(
                x=alt.X('source:N', title='Fuente', axis=alt.Axis(labelAngle=0)),
                y=alt.Y('count:Q', title='Cantidad'),
                color=alt.Color('sentiment:N', title="Sentimiento", scale=alt.Scale(
                    domain=["Positiva", "Neutra", "Negativa"],
                    range=["#2ecc71", "#f1c40f", "#e74c3c"]
                )),
                tooltip=["source", "sentiment", "count"]
            ).properties(height=400)
            st.altair_chart(chart, use_container_width=True)

    st.caption("Resúmenes con IA local (Gemma 3:4B en Ollama)")

else:
    st.warning("No se encontraron noticias para mostrar.", icon=":material/warning:")


# --------------- footer -----------------------------
st.write("---")
with st.container():
  #st.write("---")
  st.write("&copy; - derechos reservados -  2025 -  Walter Gómez - FullStack Developer - Data Science - Business Intelligence")
  #st.write("##")
  left, right = st.columns(2, gap='medium', vertical_alignment="bottom")
  with left:
    #st.write('##')
    st.link_button("Mi LinkedIn", "https://www.linkedin.com/in/walter-gomez-fullstack-developer-datascience-businessintelligence-finanzas-python/",use_container_width=True)
  with right: 
     #st.write('##') 
    st.link_button("Mi Porfolio", "https://walter-portfolio-animado.netlify.app/", use_container_width=True)
      